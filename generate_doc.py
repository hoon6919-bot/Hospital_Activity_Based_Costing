import os
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def add_heading(doc, text, level=1):
    heading = doc.add_heading(text, level=level)
    return heading

def add_paragraph(doc, text, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    if bold:
        run.bold = True
    return p

def add_bullet(doc, text):
    doc.add_paragraph(text, style='List Bullet')

def create_walkthrough():
    doc = Document()
    
    # Title
    title = doc.add_heading('병원 원가계산 시스템 - 화면별 Walkthrough', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph('본 문서는 병원 원가계산 시스템 구축 프로젝트에서 지금까지 작업된 각 화면별 구성과 주요 기능을 정리한 워크스루(Walkthrough) 자료입니다.\n')

    # 1. 대시보드
    add_heading(doc, '1. 대시보드 (DashboardPage)', level=1)
    add_bullet(doc, '주요 원가계산 현황 및 지표(수익, 비용, 손익 등)를 한눈에 파악할 수 있는 시각화 요소를 제공합니다.')
    add_bullet(doc, '연도별, 월별 데이터 비교 기능 및 추세(Trend)를 확인할 수 있습니다.')

    # 2. 기초설정
    add_heading(doc, '2. 기초설정 (SetupPage)', level=1)
    add_bullet(doc, '원가계산을 수행할 기준 연도 및 월(Period)을 설정하고 관리합니다.')
    add_bullet(doc, '부서코드, 원가대상코드 등 병원의 조직 및 원가대상 구조를 정의합니다.')
    add_bullet(doc, '시스템 전반에서 공통으로 사용되는 표준 코드(표준 활동, 표준 계정, 표준 직종, 표준 드라이버 등)를 관리합니다.')

    # 3. 비용 (회계전표)
    add_heading(doc, '3. 비용 입력 (CostPage)', level=1)
    add_bullet(doc, '병원 재무/회계 전표 데이터를 기반으로 발생한 비용 데이터를 조회하고 등록합니다.')
    add_bullet(doc, '항목별(인건비, 재료비, 관리비 등) 세부 내역을 입력하며, 다건 처리를 위한 엑셀 업로드 및 다운로드 기능을 지원합니다.')

    # 4. 사원소속 및 인건비
    add_heading(doc, '4. 사원소속 및 인건비 (LaborInputPage)', level=1)
    add_bullet(doc, '임직원의 부서 소속, 직종, 사원명 등 인적 자원 및 인건비 상세 내역을 등록합니다.')
    add_bullet(doc, '인원수 및 총 인건비를 집계하며, 엑셀 템플릿 양식 다운로드 및 데이터 업로드를 통한 일괄 등록이 가능합니다.')

    # 5. 수익
    add_heading(doc, '5. 수익 입력 (RevenueInputPage)', level=1)
    add_bullet(doc, '진료과, 시행과, 처방의사, 시행의사 등 원가대상 구조에 맞게 수익 데이터를 상세 관리합니다.')
    add_bullet(doc, '수가분류, 환자유형(외래/입원), 계정분류별 수익 내역을 입력하며, 엑셀 연동 기능을 제공합니다.')

    # 6. 환자통계
    add_heading(doc, '6. 환자통계 (PatientStatsPage)', level=1)
    add_bullet(doc, '각 부서 및 진료과/의사별로 외래, 입원 환자 수 등의 통계 지표를 관리합니다.')
    add_bullet(doc, '원가 배부의 기준(활동량)으로 사용되는 항목들을 등록하고 저장하며, 다른 화면들과 일관된 버튼 레이아웃(엑셀 처리 등)을 제공합니다.')

    # 7. 배부기준 (로직 설정)
    add_heading(doc, '7. 배부기준 (LogicPage)', level=1)
    add_bullet(doc, '비용을 배부하기 위한 단계별 규칙 및 매핑 로직을 정의하는 화면입니다.')
    add_bullet(doc, '어떤 비용을 어떤 기준(드라이버)으로 배분할지에 대한 핵심 원가계산 엔진의 규칙을 설정합니다.')

    # 8. 드라이버 설정 (드라이버 정의 및 결과)
    add_heading(doc, '8. 드라이버 설정 (AllocationValuePage)', level=1)
    doc.add_paragraph('드라이버 정의, 요약, 결과 3개의 탭으로 구성되어 원가 배부 기준값을 최종 산출하는 화면입니다.')
    add_bullet(doc, '드라이버 정의: 수익, 비용, 환자통계, 사원소속 등의 기초 데이터를 소스로 하여, 필터 조건과 합산 유형(금액, 인원수 등)을 지정하여 배부 기준(드라이버)을 생성합니다. 모달을 통해 생성된 SQL을 가독성 있게 확인할 수 있습니다.')
    add_bullet(doc, '드라이버 요약: 생성된 드라이버 데이터의 현황을 부서별, 원가대상별로 건수 및 합계로 요약하여 한눈에 파악합니다.')
    add_bullet(doc, '드라이버 결과: 드라이버별 최종 산출 값(비율, 단위 수량 등)을 상세 조회하며, 원가대상(진료과, 의사 등)을 동적으로 연결하여 사용자 친화적인 드롭박스 UI를 제공합니다. 또한 데이터 개별 삭제 기능 및 엑셀 다운로드/업로드를 지원합니다.')

    # 9. 원가계산 결과 (손익분석)
    add_heading(doc, '9. 원가계산 결과 (ResultPage)', level=1)
    add_bullet(doc, '사전 설정된 배부 로직과 드라이버 산출 값을 통해 도출된 최종 진료과별, 의사별 원가 및 손익 결과를 조회합니다.')
    add_bullet(doc, '수익과 배부된 비용을 대조하여 병원의 이익 현황과 원가 보전율 등을 시각적으로 보여줍니다.')

    # Save to docs folder
    os.makedirs('docs', exist_ok=True)
    doc.save('docs/walkthrough_all_pages.docx')
    print("Document generated successfully at docs/walkthrough_all_pages.docx")

if __name__ == '__main__':
    create_walkthrough()
